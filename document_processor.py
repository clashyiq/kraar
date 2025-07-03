# -*- coding: utf-8 -*-
"""
Teacher AI Enhanced - Document Processor
تطبيق المدرس AI المحسن - معالج المستندات

Author: Teacher AI Enhanced Team
Version: 2.0.0
"""

import os
import logging
import mimetypes
import hashlib
from pathlib import Path
from typing import Optional, Dict, List, Tuple
import tempfile
import shutil

# Document processing libraries
try:
    import PyPDF2
    from PyPDF2 import PdfReader
except ImportError:
    PyPDF2 = None
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import pytesseract
    from PIL import Image
except ImportError:
    pytesseract = None
    Image = None

# Arabic text processing
try:
    import arabic_reshaper
    from bidi.algorithm import get_display
except ImportError:
    arabic_reshaper = None
    get_display = None

# Configure logging
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Enhanced document processor with Arabic text support"""
    
    def __init__(self):
        """Initialize the document processor"""
        self.supported_types = {
            'application/pdf': self._process_pdf,
            'application/msword': self._process_doc,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'text/plain': self._process_txt,
            'application/rtf': self._process_rtf,
            'text/rtf': self._process_rtf
        }
        
        # Arabic text processing configuration
        self.arabic_config = {
            'reshape_enabled': arabic_reshaper is not None,
            'bidi_enabled': get_display is not None,
            'detect_arabic': True,
            'preserve_formatting': True
        }
        
        logger.info("✅ تم تهيئة معالج المستندات")
    
    def process_file(self, file_path: str) -> str:
        """
        Process a document file and extract text content
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"الملف غير موجود: {file_path}")
            
            # Detect MIME type
            mime_type, _ = mimetypes.guess_type(file_path)
            if not mime_type:
                mime_type = self._detect_mime_type(file_path)
            
            logger.info(f"🔍 معالجة الملف: {file_path} (النوع: {mime_type})")
            
            # Process based on file type
            if mime_type in self.supported_types:
                processor = self.supported_types[mime_type]
                content = processor(file_path)
            else:
                # Try as text file
                content = self._process_txt(file_path)
            
            # Enhance Arabic text
            if content and self._is_arabic_text(content):
                content = self._enhance_arabic_text(content)
            
            logger.info(f"✅ تم معالجة الملف بنجاح: {len(content)} حرف")
            return content
            
        except Exception as e:
            logger.error(f"❌ خطأ في معالجة الملف {file_path}: {e}")
            return f"فشل في معالجة الملف: {str(e)}"
    
    def _process_pdf(self, file_path: str) -> str:
        """Process PDF files"""
        content = ""
        
        # Try pdfplumber first (better for Arabic)
        if pdfplumber:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            content += page_text + "\n"
                logger.info("📖 تم استخدام pdfplumber لاستخراج النص")
                return content.strip()
            except Exception as e:
                logger.warning(f"⚠️ فشل pdfplumber: {e}")
        
        # Fallback to PyPDF2
        if PdfReader:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            content += page_text + "\n"
                logger.info("📖 تم استخدام PyPDF2 لاستخراج النص")
                return content.strip()
            except Exception as e:
                logger.warning(f"⚠️ فشل PyPDF2: {e}")
        
        # If both fail, try OCR for images
        if pytesseract and Image:
            try:
                return self._ocr_pdf(file_path)
            except Exception as e:
                logger.warning(f"⚠️ فشل OCR: {e}")
        
        return "فشل في استخراج النص من ملف PDF"
    
    def _process_docx(self, file_path: str) -> str:
        """Process DOCX files"""
        if not DocxDocument:
            return "مكتبة python-docx غير مثبتة"
        
        try:
            doc = DocxDocument(file_path)
            content = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content += " | ".join(row_text) + "\n"
            
            logger.info("📄 تم استخراج النص من DOCX")
            return content.strip()
            
        except Exception as e:
            logger.error(f"❌ خطأ في معالجة DOCX: {e}")
            return f"فشل في معالجة ملف DOCX: {str(e)}"
    
    def _process_doc(self, file_path: str) -> str:
        """Process DOC files (legacy Word format)"""
        # DOC files are more complex, try to convert or use textract
        try:
            # Try reading as text (might work for some DOC files)
            with open(file_path, 'rb') as file:
                content = file.read()
                # Try to decode as latin-1 and then clean up
                text = content.decode('latin-1', errors='ignore')
                # Clean up binary garbage
                cleaned_text = ''.join(char for char in text if char.isprintable() or char.isspace())
                if len(cleaned_text) > 100:  # If we got meaningful content
                    return cleaned_text
        except Exception as e:
            logger.warning(f"⚠️ فشل في قراءة DOC كنص: {e}")
        
        return "تنسيق DOC القديم يتطلب تحويل. يرجى حفظ الملف بصيغة DOCX"
    
    def _process_txt(self, file_path: str) -> str:
        """Process text files with encoding detection"""
        encodings = ['utf-8', 'utf-16', 'windows-1256', 'iso-8859-6', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                    if content and not any(ord(char) > 1000000 for char in content[:100]):
                        logger.info(f"📝 تم قراءة النص بترميز {encoding}")
                        return content.strip()
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as e:
                logger.warning(f"⚠️ خطأ في قراءة الملف بترميز {encoding}: {e}")
        
        # Last resort: read as binary and clean
        try:
            with open(file_path, 'rb') as file:
                content = file.read()
                text = content.decode('utf-8', errors='ignore')
                return text.strip()
        except Exception as e:
            logger.error(f"❌ فشل في قراءة الملف النصي: {e}")
            return "فشل في قراءة الملف النصي"
    
    def _process_rtf(self, file_path: str) -> str:
        """Process RTF files"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                
            # Simple RTF parsing - remove RTF control codes
            import re
            
            # Remove RTF header and control words
            content = re.sub(r'\\[a-z]+\d*\s*', '', content)
            content = re.sub(r'[{}]', '', content)
            content = re.sub(r'\\[^a-z]', '', content)
            
            # Clean up extra whitespace
            content = re.sub(r'\s+', ' ', content)
            
            if content.strip():
                logger.info("📄 تم استخراج النص من RTF")
                return content.strip()
            else:
                return "ملف RTF فارغ أو تالف"
                
        except Exception as e:
            logger.error(f"❌ خطأ في معالجة RTF: {e}")
            return f"فشل في معالجة ملف RTF: {str(e)}"
    
    def _ocr_pdf(self, file_path: str) -> str:
        """Extract text from PDF using OCR"""
        if not (pytesseract and Image):
            return "مكتبات OCR غير متوفرة"
        
        try:
            # Convert PDF to images and OCR
            content = ""
            # This is a simplified OCR approach
            # In production, you'd want to use pdf2image to convert PDF pages to images
            logger.info("🔍 محاولة استخراج النص باستخدام OCR")
            return "OCR غير مكتمل في هذا الإصدار"
            
        except Exception as e:
            logger.error(f"❌ خطأ في OCR: {e}")
            return f"فشل في OCR: {str(e)}"
    
    def _detect_mime_type(self, file_path: str) -> str:
        """Detect MIME type from file content"""
        try:
            with open(file_path, 'rb') as file:
                header = file.read(4096)
            
            # PDF signature
            if header.startswith(b'%PDF'):
                return 'application/pdf'
            
            # DOCX signature (ZIP-based)
            if header.startswith(b'PK\x03\x04') and b'word/' in header:
                return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            
            # DOC signature
            if header.startswith(b'\xd0\xcf\x11\xe0'):
                return 'application/msword'
            
            # RTF signature
            if header.startswith(b'{\\rtf'):
                return 'application/rtf'
            
            # Default to text
            return 'text/plain'
            
        except Exception:
            return 'text/plain'
    
    def _is_arabic_text(self, text: str) -> bool:
        """Check if text contains Arabic characters"""
        if not text:
            return False
        
        arabic_count = sum(1 for char in text if '\u0600' <= char <= '\u06FF')
        total_chars = len([char for char in text if char.isalpha()])
        
        if total_chars == 0:
            return False
        
        return arabic_count / total_chars > 0.3  # More than 30% Arabic
    
    def _enhance_arabic_text(self, text: str) -> str:
        """Enhance Arabic text for better processing"""
        if not text:
            return text
        
        try:
            # Normalize Arabic text
            enhanced_text = text
            
            # Fix common Arabic text issues
            fixes = {
                'ي': 'ي',  # Normalize Yeh
                'ك': 'ك',  # Normalize Kaf
                'ة': 'ة',  # Normalize Teh Marbuta
                'أ': 'أ',  # Normalize Alef with Hamza
                'إ': 'إ',  # Normalize Alef with Hamza below
                'آ': 'آ',  # Normalize Alef with Madda
            }
            
            for old, new in fixes.items():
                enhanced_text = enhanced_text.replace(old, new)
            
            # Reshape Arabic text if library is available
            if self.arabic_config['reshape_enabled'] and arabic_reshaper:
                try:
                    reshaped = arabic_reshaper.reshape(enhanced_text)
                    if self.arabic_config['bidi_enabled'] and get_display:
                        enhanced_text = get_display(reshaped)
                except Exception as e:
                    logger.warning(f"⚠️ فشل في تحسين النص العربي: {e}")
            
            return enhanced_text
            
        except Exception as e:
            logger.warning(f"⚠️ خطأ في تحسين النص العربي: {e}")
            return text
    
    def get_file_info(self, file_path: str) -> Dict:
        """Get detailed information about a file"""
        try:
            stat = os.stat(file_path)
            mime_type, _ = mimetypes.guess_type(file_path)
            
            if not mime_type:
                mime_type = self._detect_mime_type(file_path)
            
            # Calculate file hash
            hash_md5 = hashlib.md5()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            
            return {
                'filename': os.path.basename(file_path),
                'size': stat.st_size,
                'mime_type': mime_type,
                'created': stat.st_ctime,
                'modified': stat.st_mtime,
                'hash': hash_md5.hexdigest(),
                'supported': mime_type in self.supported_types
            }
            
        except Exception as e:
            logger.error(f"❌ خطأ في جمع معلومات الملف: {e}")
            return {}
    
    def extract_metadata(self, file_path: str) -> Dict:
        """Extract metadata from document"""
        metadata = {
            'title': '',
            'author': '',
            'subject': '',
            'creator': '',
            'producer': '',
            'creation_date': '',
            'modification_date': '',
            'language': 'ar' if self._detect_language(file_path) == 'arabic' else 'en'
        }
        
        try:
            mime_type, _ = mimetypes.guess_type(file_path)
            
            if mime_type == 'application/pdf' and PdfReader:
                # Extract PDF metadata
                with open(file_path, 'rb') as file:
                    pdf_reader = PdfReader(file)
                    if pdf_reader.metadata:
                        pdf_meta = pdf_reader.metadata
                        metadata.update({
                            'title': str(pdf_meta.get('/Title', '')),
                            'author': str(pdf_meta.get('/Author', '')),
                            'subject': str(pdf_meta.get('/Subject', '')),
                            'creator': str(pdf_meta.get('/Creator', '')),
                            'producer': str(pdf_meta.get('/Producer', '')),
                        })
            
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' and DocxDocument:
                # Extract DOCX metadata
                doc = DocxDocument(file_path)
                if doc.core_properties:
                    props = doc.core_properties
                    metadata.update({
                        'title': props.title or '',
                        'author': props.author or '',
                        'subject': props.subject or '',
                        'creator': props.author or '',
                    })
            
        except Exception as e:
            logger.warning(f"⚠️ فشل في استخراج البيانات الوصفية: {e}")
        
        return metadata
    
    def _detect_language(self, file_path: str) -> str:
        """Detect document language"""
        try:
            # Read a sample of the text
            content = self.process_file(file_path)
            if not content:
                return 'unknown'
            
            # Simple Arabic detection
            sample = content[:1000]  # First 1000 characters
            if self._is_arabic_text(sample):
                return 'arabic'
            else:
                return 'english'
                
        except Exception:
            return 'unknown'
    
    def validate_file(self, file_path: str) -> Tuple[bool, str]:
        """Validate if file can be processed"""
        try:
            if not os.path.exists(file_path):
                return False, "الملف غير موجود"
            
            if os.path.getsize(file_path) == 0:
                return False, "الملف فارغ"
            
            if os.path.getsize(file_path) > 16 * 1024 * 1024:  # 16MB
                return False, "الملف كبير جداً (الحد الأقصى 16 ميجابايت)"
            
            mime_type, _ = mimetypes.guess_type(file_path)
            if not mime_type:
                mime_type = self._detect_mime_type(file_path)
            
            if mime_type not in self.supported_types:
                return False, f"نوع الملف غير مدعوم: {mime_type}"
            
            return True, "الملف صالح للمعالجة"
            
        except Exception as e:
            return False, f"خطأ في التحقق من الملف: {str(e)}"
    
    def get_supported_types(self) -> List[str]:
        """Get list of supported file types"""
        return list(self.supported_types.keys())
    
    def get_processing_stats(self) -> Dict:
        """Get processing capabilities statistics"""
        return {
            'supported_types': len(self.supported_types),
            'pdf_support': PdfReader is not None or pdfplumber is not None,
            'docx_support': DocxDocument is not None,
            'ocr_support': pytesseract is not None and Image is not None,
            'arabic_support': self.arabic_config['reshape_enabled'],
            'libraries_available': {
                'PyPDF2': PyPDF2 is not None,
                'pdfplumber': pdfplumber is not None,
                'python-docx': DocxDocument is not None,
                'pytesseract': pytesseract is not None,
                'arabic-reshaper': arabic_reshaper is not None,
                'python-bidi': get_display is not None
            }
        }
